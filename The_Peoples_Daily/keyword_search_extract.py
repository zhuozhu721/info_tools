import PySimpleGUI as sg
from datetime import datetime, timedelta
import os
import requests
from bs4 import BeautifulSoup
from docx import Document
import shutil
import re
import PyPDF2

def get_layout_url(date_obj):
    date_str = date_obj.strftime("%Y%m/%d")
    return f"http://paper.people.com.cn/rmrb/pc/layout/{date_str}/node_01.html"

def fetch_all_layout_urls(first_layout_url):
    try:
        resp = requests.get(first_layout_url)
        resp.encoding = 'utf-8'
        soup = BeautifulSoup(resp.text, 'html.parser')
        layout_urls = set()
        for a in soup.find_all('a', href=True):
            href = a['href']
            if href.startswith('node_') and href.endswith('.html'):
                full_url = requests.compat.urljoin(first_layout_url, href)
                layout_urls.add(full_url)
        layout_urls.add(first_layout_url)
        return list(layout_urls)
    except Exception as e:
        print(f"获取版面链接失败: {e}")
        return []

def fetch_article_links_with_titles(layout_url):
    try:
        resp = requests.get(layout_url)
        resp.encoding = 'utf-8'
        soup = BeautifulSoup(resp.text, 'html.parser')
        articles = []
        for a in soup.find_all('a', href=True):
            href = a['href']
            if href.endswith('.html') and 'node_' not in href:
                full_url = requests.compat.urljoin(layout_url, href)
                title = a.text.strip()
                if title:
                    articles.append((full_url, title))
        return articles
    except Exception as e:
        print(f"获取文章链接和标题失败: {e}")
        return []

def fetch_article_content(url):
    try:
        resp = requests.get(url)
        resp.encoding = 'utf-8'
        soup = BeautifulSoup(resp.text, 'html.parser')
        content_div = soup.find('div', id='ozoom')
        paragraphs = content_div.find_all('p') if content_div else []
        content = '\n'.join([p.text.strip() for p in paragraphs])
        return content
    except Exception as e:
        print(f"获取正文失败: {e}")
        return ""

def save_to_docx(title, content, folder, date_str, layout_name):
    if layout_name.startswith("node_"):
        try:
            num = int(layout_name.replace("node_", ""))
            layout_str = f"第{num}版"
        except:
            layout_str = layout_name
    else:
        layout_str = layout_name
    safe_title = "".join(c for c in title if c.isalnum() or c in " _-")
    filename = f"{date_str}_{layout_str}_{safe_title}.docx"
    filename = os.path.join(folder, filename)
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(content)
    doc.save(filename)

# PDF下载相关函数
headers = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/79.0.3945.79 Safari/537.36"
}

def download_pdf_gui(date_str, partpath, newspaperpatch, window=None):
    today1 = date_str
    today2 = date_str[0:4] + date_str[5:7] + date_str[8:10]
    today3 = date_str.replace("-", "")
    try:
        os.makedirs(newspaperpatch, exist_ok=True)
    except:
        pass
    try:
        os.makedirs(partpath, exist_ok=True)
    except:
        shutil.rmtree(partpath)
        os.makedirs(partpath, exist_ok=True)
    filelist = os.listdir(newspaperpatch)
    pdfname = "People's.Daily.{}.pdf".format(today2)
    if pdfname in filelist:
        if window:
            window['-STATUS-'].update("该日期已经下载过了!")
        return
    coverurl = f"http://paper.people.com.cn/rmrb/html/{today1}/nbs.D110000renmrb_01.htm"
    response = requests.get(coverurl, headers=headers)
    pagenum = len(re.findall("nbs", response.text))
    if pagenum != 0:
        if response.status_code == 403:
            if window:
                window['-STATUS-'].update("你选择的日期太久远，网站不提供。只有两年之内的。")
            return
        for page in range(1, pagenum + 1):
            success = False
            for retry in range(2):  # 最多重试2次
                downtplurl = "http://paper.people.com.cn/rmrb/images/{0}/{2}/rmrb{1}{2}.pdf"
                formatpage = "{0:0>2}".format(page)
                downurl = downtplurl.format(today1, today2, formatpage)
                filename = 'rmrb{}.pdf'.format(today2 + formatpage)
                try:
                    response = requests.get(downurl, headers=headers, timeout=10)
                    file = response.content
                    if len(file) > 1000:
                        with open(os.path.join(partpath, filename), "wb") as fn:
                            fn.write(file)
                        success = True
                        break
                except Exception:
                    continue
            if window and page % 2 == 0:
                window['-STATUS-'].update(f"第{page}页下载中……")
        if window:
            window['-STATUS-'].update(f"{today1} PDF各版下载完成，正在合并...")
    else:
        coverurl = f"http://paper.people.com.cn/rmrb/pc/layout/{today3}/node_01.html"
        response = requests.get(coverurl, headers=headers)
        pagenum = len(re.findall("pageLink", response.text))
        for page in range(1, pagenum + 1):
            success = False
            for retry in range(2):
                currentPageUrl = f"http://paper.people.com.cn/rmrb/pc/layout/{today3}/node_{page:02d}.html"
                try:
                    response = requests.get(currentPageUrl, headers=headers, timeout=10)
                    dumpUrls = re.findall(r'''attachement.*?\.pdf''', response.text)
                    if not dumpUrls:
                        continue
                    dumpUrl = dumpUrls[0]
                    downloadUrl = "http://paper.people.com.cn/rmrb/pc/" + dumpUrl
                    formatpage = "{0:0>2}".format(page)
                    filename = 'rmrb{}.pdf'.format(today2 + formatpage)
                    response = requests.get(downloadUrl, headers=headers, timeout=10)
                    file = response.content
                    if len(file) > 1000:
                        with open(os.path.join(partpath, filename), "wb") as fn:
                            fn.write(file)
                        success = True
                        break
                except Exception:
                    continue
            if window and page % 2 == 0:
                window['-STATUS-'].update(f"第{page}页下载中……")
        if window:
            window['-STATUS-'].update(f"{today1} PDF各版下载完成，正在合并...")

def merge_pdf(partpath, newspaperpatch):
    filelist = os.listdir(partpath)
    filelist.sort()
    try:
        pdfFM = PyPDF2.PdfFileMerger(strict=False)
    except:
        pdfFM = PyPDF2.PdfMerger(strict=False)
    for file in filelist:
        fullpath = partpath + '/' + file
        filesize = os.path.getsize(fullpath)
        if filesize < 10:
            continue
        pdfFM.append(fullpath)
    pdfFM.write(newspaperpatch + "/People's.Daily." + filelist[0][4:12] + ".pdf")
    pdfFM.close()

def delete_part(partpath):
    shutil.rmtree(partpath)

def main_gui():
    today = datetime.now()
    week_ago = today - timedelta(days=6)
    default_folder = os.path.join(os.path.expanduser("~"), "Desktop")
    layout = [
        [sg.Text('请选择功能：')],
        [sg.Radio('关键词筛选下载', 'MODE', default=True, key='mode_word', enable_events=True),
         sg.Radio('整版PDF下载', 'MODE', key='mode_pdf', enable_events=True)],
        [sg.Text('起始日期(YYYY-MM-DD):'), sg.Input(default_text=week_ago.strftime("%Y-%m-%d"), key='start')],
        [sg.Text('结束日期(YYYY-MM-DD):'), sg.Input(default_text=today.strftime("%Y-%m-%d"), key='end')],
        [sg.Text('关键词(例：人工智能，科技):'), sg.Input(key='keywords')],
        [sg.Text('保存文件夹:'), sg.Input(default_text=default_folder, key='folder'), sg.FolderBrowse()],
        [sg.Text('', size=(60,1), key='-STATUS-')],
        [sg.ProgressBar(100, orientation='h', size=(50, 20), key='-PROGRESS-')],
        [
            sg.Column([
                [sg.Output(size=(60, 15))]
            ], vertical_alignment='top'),
            sg.Column([
                [sg.Multiline(
                    '使用提示：\n'
                    '1. 请选择功能。\n'
                    '2. 填写起始日期、结束日期。\n'
                    '3. “关键词筛选下载”时填写关键词。\n'
                    '4. 选择保存文件夹。\n'
                    '5. 填写完毕后点击“开始”按钮。\n',
                    size=(32, 15), disabled=True, no_scrollbar=True, border_width=0,
                    text_color='#444444', background_color=sg.theme_background_color()
                )]
            ], vertical_alignment='top', pad=((10, 0), 0))
        ],
        [sg.Button('开始'), sg.Button('退出')],
        [sg.Text('ZZ', font=('Arial', 10), text_color='#CCCCCC', justification='right', pad=((0, 10), (10, 0)), expand_x=True)]
    ]
    window = sg.Window('人民日报工具', layout, finalize=True)

    # 初始时关键词输入框可用
    window['keywords'].update(disabled=False)

    while True:
        event, values = window.read()
        if event in (sg.WINDOW_CLOSED, '退出'):
            break

        # 根据功能选择动态禁用/启用关键词输入框
        if event in ('mode_word', 'mode_pdf'):
            window['keywords'].update(disabled=values['mode_pdf'])

        if event == '开始':
            start_input = values['start'].strip()
            end_input = values['end'].strip()
            folder = values['folder'].strip()
            if not folder:
                sg.popup_error("请选择保存文件夹！")
                continue
            try:
                if not start_input:
                    start_date = today
                else:
                    start_date = datetime.strptime(start_input, "%Y-%m-%d")
                if not end_input:
                    end_date = today
                else:
                    end_date = datetime.strptime(end_input, "%Y-%m-%d")
            except Exception as e:
                print("日期格式错误，请输入YYYY-MM-DD格式")
                continue

            if start_date > end_date:
                print("起始日期不能晚于结束日期！")
                continue

            date_list = [start_date + timedelta(n) for n in range((end_date - start_date).days + 1)]

            if values['mode_word']:
                # 关键词筛选下载Word
                keywords_input = values['keywords'].strip()
                raw_keywords = keywords_input.replace('，', ',')
                keywords = [k.strip() for k in raw_keywords.split(",") if k.strip()]
                if not keywords:
                    sg.popup_error("请输入至少一个关键词！")
                    continue
                if start_date == end_date:
                    save_folder = os.path.join(folder, start_date.strftime("%Y-%m-%d"))
                else:
                    save_folder = os.path.join(folder, f"{start_date.strftime('%Y-%m-%d')}_to_{end_date.strftime('%Y-%m-%d')}")
                os.makedirs(save_folder, exist_ok=True)

                # 收集所有文章链接和标题（带进度）
                window['-STATUS-'].update("正在收集文章链接 ...")
                total_days = len(date_list)
                window['-PROGRESS-'].update_bar(0, total_days)
                window.refresh()
                article_info_list = []
                for day_idx, single_date in enumerate(date_list, 1):
                    today_str = single_date.strftime("%Y-%m-%d")
                    window['-STATUS-'].update(f"正在收集 {today_str} 的文章链接 ({day_idx}/{total_days}) ...")
                    window['-PROGRESS-'].update_bar(day_idx)
                    window.refresh()
                    first_layout_url = get_layout_url(single_date)
                    layout_urls = fetch_all_layout_urls(first_layout_url)
                    for layout_url in layout_urls:
                        layout_name = layout_url.split('/')[-1].replace('.html', '')
                        articles = fetch_article_links_with_titles(layout_url)
                        for url, title in articles:
                            article_info_list.append((url, single_date.strftime("%Y-%m-%d"), layout_name, title))
                total_articles = len(article_info_list)
                if total_articles == 0:
                    window['-STATUS-'].update("未找到任何文章，请检查日期范围。")
                    continue

                window['-PROGRESS-'].update_bar(0, total_articles)
                window['-STATUS-'].update(f"共需筛选 {total_articles} 篇文章，开始筛选 ...")
                window.refresh()

                # 只下载命中关键词的文章
                matched_count = 0
                for idx, (url, date_str, layout_name, title) in enumerate(article_info_list, 1):
                    hit_keywords = [kw for kw in keywords if kw in title]
                    if hit_keywords:
                        content = fetch_article_content(url)
                        for kw in hit_keywords:
                            # 为每个命中关键词建立子文件夹
                            kw_folder = os.path.join(save_folder, kw)
                            os.makedirs(kw_folder, exist_ok=True)
                            save_to_docx(title, content, kw_folder, date_str, layout_name)
                        matched_count += 1
                        print(f"[{matched_count}] {date_str} 命中: {title}（关键词：{'，'.join(hit_keywords)}）")
                    window['-PROGRESS-'].update_bar(idx)
                    if idx % 10 == 0 or idx == total_articles:
                        window['-STATUS-'].update(f"已筛选 {idx}/{total_articles} 篇，命中 {matched_count} 篇 ...")
                        window.refresh()
                window['-STATUS-'].update(f"全部完成！共保存 {matched_count} 篇命中文章。")
                print(f"全部完成！共保存 {matched_count} 篇命中文章。")

            elif values['mode_pdf']:
                # 整版PDF下载
                if start_date == end_date:
                    save_folder = os.path.join(folder, start_date.strftime("%Y-%m-%d") + "_PDF")
                else:
                    save_folder = os.path.join(folder, f"{start_date.strftime('%Y-%m-%d')}_to_{end_date.strftime('%Y-%m-%d')}_PDF")
                os.makedirs(save_folder, exist_ok=True)
                total_days = len(date_list)
                window['-PROGRESS-'].update_bar(0, total_days)
                for day_idx, single_date in enumerate(date_list, 1):
                    window['-STATUS-'].update(f"正在下载 {single_date.strftime('%Y-%m-%d')} 的PDF ({day_idx}/{total_days}) ...")
                    window['-PROGRESS-'].update_bar(day_idx)
                    window.refresh()
                    # PDF下载到临时part文件夹
                    partpath = os.path.join(save_folder, "part")
                    newspaperpatch = save_folder
                    download_pdf_gui(single_date.strftime("%Y-%m/%d"), partpath, newspaperpatch, window)
                    merge_pdf(partpath, newspaperpatch)
                    delete_part(partpath)
                    window['-STATUS-'].update(f"{single_date.strftime('%Y-%m-%d')} PDF下载并合并完成。")
                window['-STATUS-'].update("PDF下载全部完成！")
                print("PDF下载全部完成！")

    window.close()

if __name__ == "__main__":
    main_gui()