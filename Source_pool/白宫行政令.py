from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from bs4 import BeautifulSoup
import os
import time
from datetime import datetime

def collect(start_dt, end_dt, save_folder=None):
    results = []
    chromedriver_path = r"C:\Program Files\Google\Chrome_107.0.5304.122\Chrome-bin\chromedriver_win32\chromedriver.exe"
    chrome_path = r"C:\Program Files\Google\Chrome_107.0.5304.122\Chrome-bin\chrome.exe"
    chrome_options = Options()
    chrome_options.add_argument("--headless")
    chrome_options.add_argument("--disable-gpu")
    chrome_options.add_argument("--no-sandbox")
    chrome_options.binary_location = chrome_path
    service = Service(executable_path=chromedriver_path)
    driver = webdriver.Chrome(service=service, options=chrome_options)
    base_url = "https://www.whitehouse.gov"
    list_url = "https://www.whitehouse.gov/presidential-actions/executive-orders/"
    page = 1
    while True:
        url = list_url if page == 1 else f"{list_url}page/{page}/"
        driver.get(url)
        time.sleep(3)
        driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
        time.sleep(3)
        html = driver.page_source
        soup = BeautifulSoup(html, "html.parser")
        blocks = soup.select('div.wp-block-group.wp-block-whitehouse-post-template__content')
        has_article_in_range = False
        for block in blocks:
            h2 = block.select_one('h2.wp-block-post-title')
            a = h2.find('a') if h2 else None
            title = a.get_text(strip=True) if a else ""
            href = a['href'] if a else ""
            if not href.startswith("http"):
                href = base_url + href
            category = ""
            cat_tag = block.select_one('.wp-block-post-terms a')
            if cat_tag:
                category = cat_tag.get_text(strip=True)
            else:
                cat_text = block.get_text(separator="|", strip=True)
                if "Executive Orders" in cat_text:
                    category = "Executive Orders"
            date_str = ""
            date_tag = block.select_one('time')
            if date_tag:
                date_str = date_tag.get_text(strip=True)
                try:
                    dt = datetime.strptime(date_str, "%B %d, %Y")
                except Exception:
                    try:
                        dt = datetime.strptime(date_str, "%b %d, %Y")
                    except Exception:
                        dt = None
            else:
                dt = None
            if not dt or not (start_dt <= dt <= end_dt):
                continue
            has_article_in_range = True
            summary = ""
            try:
                driver.get(href)
                time.sleep(2)
                detail_html = driver.page_source
                detail_soup = BeautifulSoup(detail_html, "html.parser")
                content = (
                    detail_soup.select_one(".wp-block-post-content") or
                    detail_soup.select_one(".page-content__content") or
                    detail_soup.select_one(".body-content")
                )
                if content:
                    text = content.get_text(separator="\n", strip=True)
                    summary = text[:200].replace('\n', ' ')
                    # 保存正文到分信息源子目录，文件名带来源和日期，支持外部传入保存目录
                    try:
                        import os
                        from docx import Document
                        from docx.oxml.ns import qn
                        from docx.shared import Pt
                        save_dir = None
                        import inspect
                        frame = inspect.currentframe()
                        while frame:
                            if 'save_folder' in frame.f_locals:
                                save_folder = frame.f_locals['save_folder']
                                if save_folder:
                                    save_dir = os.path.join(save_folder, '白宫行政令')
                                break
                            frame = frame.f_back
                        if not save_dir:
                            save_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '../../download/白宫行政令'))
                        if not os.path.exists(save_dir):
                            os.makedirs(save_dir)
                        doc = Document()
                        # 标题样式
                        heading = doc.add_heading(level=1)
                        run = heading.add_run(title)
                        run.font.name = u'方正小标宋简体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'方正小标宋简体')
                        run.font.size = Pt(16)
                        # 正文样式
                        style = doc.styles['Normal']
                        font = style.font
                        font.name = u'仿宋_GB2312'
                        font.size = Pt(12)
                        style.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
                        doc.add_paragraph(text)
                        safe_title = "".join([c if c.isalnum() else "_" for c in title])[:50]
                        file_name = f"白宫行政令_{date_str}_{safe_title}.docx"
                        doc.save(os.path.join(save_dir, file_name))
                        print(f"已保存: {file_name}")
                    except Exception as e:
                        print(f"保存docx异常: {e}")
            except Exception:
                summary = ""
            results.append({
                "date": date_str,
                "title": title,
                "url": href,
                "source": "白宫行政令",
                "category": category,
                "summary": summary
            })
        if not has_article_in_range or len(blocks) == 0:
            break
        page += 1
    driver.quit()
    return results