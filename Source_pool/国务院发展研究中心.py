import requests
from datetime import datetime
import time

def collect(start_dt, end_dt):
    results = []
    url = "https://www.drc.gov.cn/Json/GetPageDocuments.ashx"
    headers = {
        "User-Agent": "Mozilla/5.0",
        "Referer": "https://www.drc.gov.cn/Leaf.aspx?leafid=1338",
        "Accept": "application/json, text/javascript, */*; q=0.01",
        "Accept-Language": "zh-CN,zh;q=0.9"
    }
    page = 1
    while True:
        params = {
            "chnid": 378,
            "leafid": 1338,
            "page": page,
            "pagesize": 10,
            "sublen": 21,
            "sumlen": 230,
            "keyword": "",
            "expertid": 0
        }
        try:
            response = requests.get(url, params=params, headers=headers, timeout=10)
            if response.status_code != 200:
                break
            data = response.json()
            #print(data)
        except Exception as e:
            print(f"国务院发展研究中心第{page}页请求失败: {e}")
            break

        if not (isinstance(data, list) and len(data) > 0 and "rows" in data[0]):
            break

        for item in data[0]["rows"]:
            date_str = item.get("DelivedDate", "")[:10]
            title = item.get("Subject", "")
            url_detail = item.get("DocViewUrl", "")
            print(f"title: {title} | DocViewUrl: {url_detail}")
            # 自动补全相对路径
            if url_detail and not url_detail.startswith("http"):
                url_detail = "https://www.drc.gov.cn/" + url_detail.lstrip("/")

            # 下载正文并保存为docx
            content = ""
            try:
                from bs4 import BeautifulSoup
                from docx import Document
                from docx.oxml.ns import qn
                from docx.shared import Pt
            except ImportError:
                print("缺少依赖库，请先安装 beautifulsoup4 和 python-docx")
                continue

            if url_detail and url_detail.startswith("http"):
                try:
                    detail_resp = requests.get(url_detail, headers=headers, timeout=10)
                    #print(f"请求详情页: {url_detail}，状态码: {detail_resp.status_code}，实际URL: {detail_resp.url}")
                    #print(f"响应内容前200字: {detail_resp.text[:200]}")
                    detail_resp.encoding = detail_resp.apparent_encoding
                    soup = BeautifulSoup(detail_resp.text, "html.parser")
                    main_content = soup.find("div", id="MainContent_docContent")
                    if main_content:
                        # 保留段落换行
                        content = "\n".join(p.get_text(strip=True) for p in main_content.find_all("p"))
                    else:
                        content = ""
                except Exception as e:
                    print(f"详情页请求异常: {e}")
                    content = ""
            if content:
                # 保存为docx，分信息源子目录，文件名带来源和日期
                try:
                    import os
                    save_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '../../download/国务院发展研究中心'))
                    if not os.path.exists(save_dir):
                        os.makedirs(save_dir)
                    doc = Document()
                    # 正文样式
                    style = doc.styles['Normal']
                    font = style.font
                    font.name = u'仿宋_GB2312'
                    font.size = Pt(12)
                    style.element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
                    # 标题样式
                    heading = doc.add_heading(level=1)
                    run = heading.add_run(title)
                    run.font.name = u'方正小标宋简体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'方正小标宋简体')
                    run.font.size = Pt(16)
                    doc.add_paragraph(content)
                    safe_title = "".join([c if c.isalnum() else "_" for c in title])[:50]
                    file_name = f"国务院发展研究中心_{date_str}_{safe_title}.docx"
                    doc.save(os.path.join(save_dir, file_name))
                    print(f"已保存: {file_name}")
                except Exception as e:
                    print(f"保存docx异常: {e}")

            # 已彻底移除冗余的 drc_articles 保存逻辑，仅保留 download 目录保存
            try:
                dt = datetime.strptime(date_str, "%Y-%m-%d")
            except Exception:
                continue
            if start_dt <= dt <= end_dt:
                results.append({
                    "date": date_str,
                    "title": title,
                    "url": url_detail,
                    "source": "国务院发展研究中心",
                    "category": "",
                    "summary": ""
                })

        all_earlier = all(
            datetime.strptime(item.get("DelivedDate", "")[:10], "%Y-%m-%d") < start_dt
            for item in data[0]["rows"] if item.get("DelivedDate", "")
        )
        if all_earlier:
            break

        page += 1
        time.sleep(1)

    return results