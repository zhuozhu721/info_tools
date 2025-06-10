import os
import re
import json
import requests
from datetime import datetime
from bs4 import BeautifulSoup

def parse_preloaded_state(html):
    # 提取 window.__PRELOADED_STATE__ = {...};
    m = re.search(r'window\.__PRELOADED_STATE__\s*=\s*(\{.*?\});?\s*</script>', html, re.DOTALL)
    if not m:
        return None
    try:
        data = json.loads(m.group(1))
        return data
    except Exception:
        return None

def get_most_popular_list(preloaded_state):
    # Most Popular 区块在 homepage-section 03，gallery-section，title: Most Popular
    # 新版MIT首页已无Most Popular区块在preloaded_state，直接用BeautifulSoup抓取首页HTML
    try:
        import requests
        from bs4 import BeautifulSoup
        base_url = 'https://www.technologyreview.com'
        resp = requests.get(base_url, timeout=10)
        if resp.status_code != 200:
            return []
        soup = BeautifulSoup(resp.text, 'html.parser')
        # Most Popular区块通常有明显的标题
        most_popular_header = soup.find(lambda tag: tag.name in ['h2', 'h3'] and 'Most Popular' in tag.get_text())
        if not most_popular_header:
            print("[MIT_Tech] 未找到Most Popular区块标题")
            return []
        # 继续向上找，直到找到包含新闻列表的gallerySection父节点
        parent = most_popular_header
        for _ in range(5):
            if parent and parent.name == 'section' and 'gallerySection' in (parent.get('class') or ['']):
                break
            parent = parent.parent
        # 若没找到section，尝试找最近的section或div
        if not (parent and parent.name in ['section', 'div']):
            parent = most_popular_header.parent
        # 查找所有新闻条目（通常是a标签或li/article等）
        items = []
        # 先找所有article，且正文长度大于一定阈值才算文章
        for article in parent.find_all('article'):
            a = article.find('a', href=True)
            if a:
                title = a.get_text(strip=True)
                url = a['href']
                if url and not url.startswith('http'):
                    url = base_url + url
                # 进一步判断：正文长度大于200字才算文章
                content, _ = fetch_article_content(url)
                if title and len(title) > 10 and len(content) > 200:
                    items.append({'title': title, 'url': url})
            if len(items) >= 10:
                break
        # 如果没找到，降级找a标签
        if not items:
            for link in parent.find_all('a', href=True):
                title = link.get_text(strip=True)
                url = link['href']
                if url and not url.startswith('http'):
                    url = base_url + url
                content, _ = fetch_article_content(url)
                if title and len(title) > 10 and len(content) > 200:
                    items.append({'title': title, 'url': url})
                if len(items) >= 10:
                    break
        print(f"[MIT_Tech] Most Popular区块采集到{len(items)}条")
        return items
    except Exception as e:
        print(f"[MIT_Tech] get_most_popular_list异常: {e}")
        return []

def fetch_article_content(url):
    try:
        resp = requests.get(url, timeout=10)
        if resp.status_code != 200:
            return '', ''
        soup = BeautifulSoup(resp.text, 'html.parser')
        # 取正文：优先找article、section、div等含有新闻内容的块，避免header/footer/宣传
        main = soup.find('main')
        content = ''
        if main:
            article = main.find('article')
            if article:
                paragraphs = article.find_all('p')
                content = '\n'.join([p.get_text(strip=True) for p in paragraphs if p.get_text(strip=True)])
            else:
                paragraphs = [p for p in main.find_all('p') if not p.find_parent(['header', 'footer', 'aside'])]
                content = '\n'.join([p.get_text(strip=True) for p in paragraphs if p.get_text(strip=True)])
        if not content:
            article = soup.find('article')
            if article:
                paragraphs = article.find_all('p')
                content = '\n'.join([p.get_text(strip=True) for p in paragraphs if p.get_text(strip=True)])
        if not content:
            paragraphs = [p for p in soup.find_all('p') if not p.find_parent(['header', 'footer', 'aside'])]
            content = '\n'.join([p.get_text(strip=True) for p in paragraphs if p.get_text(strip=True)])

        # 进一步：通过源码结构类别(class/id)过滤广告/推荐/订阅等无关内容
        # 先用原有tail_patterns做一次截断
        tail_patterns = [
            r"Share$", r"Popular$", r"Deep Dive$", r"Related Story$", r"^\s*Subscribe", r"^\s*Sign up", r"^\s*Read next", r"^\s*Read more", r"^\s*Recommended", r"^\s*More from", r"^\s*MIT Technology Review", r"^\s*The latest iteration of a legacy", r"^\s*Climate change and energy", r"^\s*Advertise with MIT Technology Review",
            r"^Discover special offers", r"^Thank you for submitting your email", r"^It looks like something went wrong", r"^We’re having trouble saving your preferences", r"^Try refreshing this page", r"^reach out to us at", r"^upcoming events", r"^If you continue to get this message"
        ]
        lines = content.split('\n')
        def is_tail(line):
            for pat in tail_patterns:
                if re.search(pat, line, re.IGNORECASE):
                    return True
            return False
        for i in range(len(lines)-1, -1, -1):
            if not is_tail(lines[i]) and lines[i].strip():
                lines = lines[:i+1]
                break
        content = '\n'.join(lines).strip()

        # 用BeautifulSoup再次过滤：去除常见广告/推荐/订阅等区块（如class/id包含ad, promo, subscribe, newsletter, recommended, offer, footer, sidebar, modal, pop, cookie等）
        def is_junk_tag(tag):
            if not tag.name:
                return False
            attrs = ' '.join([str(v) for v in tag.attrs.values()])
            junk_keywords = [
                'ad', 'promo', 'subscribe', 'newsletter', 'recommended', 'offer', 'footer', 'sidebar', 'modal', 'pop', 'cookie', 'related', 'social', 'share', 'disclaimer', 'consent', 'alert', 'banner', 'notice', 'marketing', 'survey', 'signup', 'register', 'author-info', 'authorInfo', 'byline', 'nav', 'breadcrumb', 'utility', 'widget', 'sponsor', 'trending', 'popular', 'deep-dive', 'deepDive', 'read-next', 'readmore', 'read-more', 'readNext', 'readMore', 'thank-you', 'thankyou', 'error', 'success', 'fail', 'feedback', 'contact', 'help', 'support', 'login', 'logout', 'user', 'profile', 'avatar', 'comment', 'reply', 'form', 'input', 'search', 'filter', 'sort', 'pagination', 'page-nav', 'pageNav', 'page-navigation', 'pageNavigation', 'tab', 'tabs', 'panel', 'accordion', 'collapse', 'expand', 'toggle', 'dropdown', 'menu', 'list', 'item', 'icon', 'image', 'img', 'video', 'audio', 'media', 'embed', 'iframe', 'object', 'script', 'style', 'link', 'meta', 'head', 'body', 'html', 'main', 'section', 'article', 'aside', 'div', 'span', 'p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6'
            ]
            for kw in junk_keywords:
                if kw in attrs.lower():
                    return True
            return False

        # 只保留正文主要内容块（article/main下的p，且父节点class/id不含上述关键词）
        main = soup.find('main')
        article = main.find('article') if main else soup.find('article')
        if article:
            clean_paragraphs = []
            for p in article.find_all('p'):
                # 只保留父节点class/id不含junk关键词的段落
                parent = p.parent
                is_junk = False
                while parent and parent != article:
                    if is_junk_tag(parent):
                        is_junk = True
                        break
                    parent = parent.parent
                if not is_junk:
                    txt = p.get_text(strip=True)
                    if txt:
                        clean_paragraphs.append(txt)
            if clean_paragraphs:
                content = '\n'.join(clean_paragraphs)
        # 取发布日期（多策略）
        date = ''
        # 1. meta property="article:published_time"
        date_tag = soup.find('meta', {'property': 'article:published_time'})
        if date_tag and date_tag.get('content'):
            date_str = date_tag['content']
            try:
                date = datetime.fromisoformat(date_str.replace('Z', '+00:00')).strftime('%Y-%m-%d')
            except Exception:
                date = date_str[:10]
        # 2. meta name="date" 或 name="pubdate"
        if not date:
            for key in ['date', 'pubdate', 'publishdate', 'dc.date']:
                tag = soup.find('meta', {'name': key})
                if tag and tag.get('content'):
                    date = tag['content'][:10]
                    break
        # 3. time标签
        if not date:
            t = soup.find('time')
            if t and t.get('datetime'):
                date = t['datetime'][:10]
            elif t and t.get_text(strip=True):
                date = t.get_text(strip=True)[:10]
        # 4. 文章头部可能有日期
        if not date:
            m = re.search(r'(20\d{2}-\d{2}-\d{2})', soup.get_text())
            if m:
                date = m.group(1)
        # 5. 从URL中提取日期（如/2025/05/01/）
        if not date:
            m = re.search(r'/((20\d{2})/(\d{2})/(\d{2}))/', url)
            if m:
                date = f"{m.group(2)}-{m.group(3)}-{m.group(4)}"
        return content, date
    except Exception:
        return '', ''

def collect(start_dt=None, end_dt=None, keywords=None, save_folder=None):
    source_name = 'MIT科技评论10大热点'
    base_url = 'https://www.technologyreview.com'
    # 保证MIT TECH文章单独存放在 download/MIT科技评论(MIT Tech Review) 目录下
    if save_folder:
        save_folder = os.path.join(save_folder, source_name)
    else:
        save_folder = os.path.join(os.path.expanduser('~'), 'Desktop', 'download', source_name)
    os.makedirs(save_folder, exist_ok=True)
    resp = requests.get(base_url, timeout=10)
    if resp.status_code != 200:
        return []
    preloaded_state = parse_preloaded_state(resp.text)
    if preloaded_state is None:
        return []
    posts = get_most_popular_list(preloaded_state)
    print(f"[MIT_Tech] Most Popular新闻数: {len(posts)}")
    results = []
    from docx import Document
    for idx, post in enumerate(posts[:10], 1):
        try:
            # 新结构：post为{'title':..., 'url':...}
            title = post.get('title', '')
            url = post.get('url', '')
            summary = ''
            topic = ''
            #print(f"[MIT_Tech] ({idx}) 标题: {title}")
            # 抓正文和日期
            content, date = fetch_article_content(url)
            #print(f"[MIT_Tech] ({idx}) 日期: {date}，正文长度: {len(content)}")
            # 不再限制日期范围，top10全部保存
            item = {
                'title': title,
                'date': date,
                'url': url,
                'source': source_name,
                'summary': summary,
                'topic': topic,
                'content': content
            }
            # 保存为docx，文件名带日期
            # 文件名和标题都强制带日期
            file_date = date or ''
            # 如果date为空，尝试再次从url提取
            if not file_date:
                m = re.search(r'/((20\d{2})/(\d{2})/(\d{2}))/', url)
                if m:
                    file_date = f"{m.group(2)}-{m.group(3)}-{m.group(4)}"
            # 用完整标题，去除特殊字符，避免文件名过长或非法
            safe_title = re.sub(r'[\\/:*?"<>|]', '_', title)
            fname = f"MIT_Tech_{file_date or 'nodate'}_{safe_title}.docx"
            fpath = os.path.join(save_folder, fname)
            doc = Document()
            doc.add_heading(f"[{file_date or 'nodate'}] {title}", 0)
            doc.add_paragraph(f"Date: {file_date}")
            doc.add_paragraph(f"URL: {url}")
            doc.add_paragraph(f"Source: {source_name}")
            doc.add_paragraph(f"Topic: {topic}")
            doc.add_paragraph(f"Summary: {summary}")
            doc.add_paragraph("")
            # 美化正文：每段分行，段落间加空行
            for para in content.split('\n'):
                p = para.strip()
                if p:
                    doc.add_paragraph(p)
                    doc.add_paragraph("")
            doc.save(fpath)
            print(f"[MIT_Tech] ({idx}) 已保存: {fpath}")
            results.append(item)
        except Exception as e:
            print(f"[MIT_Tech] ({idx}) 采集异常: {e}")
            continue
    print(f"[MIT_Tech] 返回{len(results)}条新闻")
    return results

if __name__ == '__main__':
    # 测试采集
    from datetime import datetime, timedelta
    today = datetime.now()
    week_ago = today - timedelta(days=7)
    collect(start_dt=week_ago, end_dt=today)
